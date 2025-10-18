#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import os
import re
import sys
from dataclasses import dataclass, field
from typing import List, Optional, Dict, Tuple

from docx import Document
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn

# -----------------------
# Normalización y claves
# -----------------------

def norm(s: str) -> str:
    return (s or "").replace("\xa0", " ").strip()

SECTION_KEYS = {
    "PREGUNTA": "pregunta",
    "COMPETENCIA": "competencia",
    "EVIDENCIA": "evidencia",
    "CONTENIDO": "contenido",
    "CONTEXTO": "contexto",
    "RESPUESTA CORRECTA": "respuesta_correcta",
    "ENUNCIADO": "enunciado",
    "OPCIONES DE RESPUESTA": "opciones",
    "AYUDA 1": "ayuda_1",
    "AYUDA 2": "ayuda_2",
}

OPTION_LABELS = ["A", "B", "C", "D", "E", "F"]

def header_key(text: str) -> Optional[str]:
    t = norm(text)
    t = re.sub(r"\s+", " ", t).upper().rstrip(":")
    if re.fullmatch(r"AYUDA\s*1", t): return SECTION_KEYS["AYUDA 1"]
    if re.fullmatch(r"AYUDA\s*2", t): return SECTION_KEYS["AYUDA 2"]
    for k in SECTION_KEYS.keys():
        if t == k:
            return SECTION_KEYS[k]
    return None

def looks_like_question_number(text: str) -> Optional[int]:
    t = norm(text)
    m = re.match(r"^(?:N[º°]\s*)?(\d{1,4})\.?$", t)
    if m:
        return int(m.group(1))
    return None

def sanitize_bullets(text: str) -> str:
    t = text or ""
    t = re.sub(r"^[\u2022\u00B7\u25CF\u2023\u2043\u2219\u25A0\u25AA\u25CB\u25C9\u25D8\u25E6\u25E7\u25D9\u25C8\uF0B7]\s*", "- ", t)
    t = t.replace("	", "    ")
    return t

# -----------------------
# Iteración en orden
# -----------------------

def iter_block_items(doc: Document):
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P

    parent_elm = doc.element.body
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield _paragraph_from_xml(doc, child)
        elif isinstance(child, CT_Tbl):
            yield _table_from_xml_element(doc, child)

def _paragraph_from_xml(document: Document, p_elm):
    for p in document.paragraphs:
        if p._element is p_elm:
            return p
    return None

def _table_from_xml_element(document: Document, tbl_elm):
    for t in document.tables:
        if t._element is tbl_elm:
            return t
    return None

# -----------------------
# Imágenes (XPath sin namespaces)
# -----------------------

def extract_images_from_paragraph(paragraph: Paragraph) -> List[Tuple[bytes, str]]:
    images: List[Tuple[bytes, str]] = []
    if paragraph is None:
        return images
    pics = paragraph._element.xpath('.//*[local-name()="pic"]')
    for pic in pics:
        blips = pic.xpath('.//*[local-name()="blip"]')
        if not blips:
            continue
        rEmbed = blips[0].get(qn("r:embed"))
        if not rEmbed:
            continue
        image_part = paragraph.part.related_parts.get(rEmbed)
        if not image_part:
            continue
        blob = image_part.blob
        ext = os.path.splitext(image_part.filename)[1].lower() or ".png"
        images.append((blob, ext))
    return images

def _extract_images_from_cell_recursive(cell: _Cell, acc: List[Tuple[bytes, str]]):
    for p in cell.paragraphs:
        acc.extend(extract_images_from_paragraph(p))
    for t in getattr(cell, "tables", []):
        for r in t.rows:
            for c in r.cells:
                _extract_images_from_cell_recursive(c, acc)

def extract_images_from_cell(cell: _Cell) -> List[Tuple[bytes, str]]:
    out: List[Tuple[bytes, str]] = []
    _extract_images_from_cell_recursive(cell, out)
    return out

# -----------------------
# Tablas -> Markdown
# -----------------------

def cell_text(cell: _Cell) -> str:
    texts = [norm(p.text) for p in cell.paragraphs]
    texts = [t for t in texts if t]
    return "\n".join(texts)

def nested_tables_to_md(cell: _Cell) -> str:
    parts: List[str] = []
    for t in getattr(cell, "tables", []):
        parts.append(table_to_markdown(t))
    return "\n\n".join([p for p in parts if p])

def table_to_markdown(table: Table) -> str:
    data = []
    max_cols = 0
    for row in table.rows:
        row_vals = [cell_text(c) for c in row.cells]
        data.append(row_vals)
        max_cols = max(max_cols, len(row_vals))
    if not data:
        return ""
    data = [row + [""]*(max_cols - len(row)) for row in data]
    header = data[0]
    lines = []
    lines.append("| " + " | ".join(header) + " |")
    lines.append("|" + "|".join(["---"]*max_cols) + "|")
    for row in data[1:]:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)

# -----------------------
# Modelos y helpers
# -----------------------

@dataclass
class Opcion:
    id: str
    tipo: str  # "texto" | "imagen"
    texto: Optional[str] = None
    imagen: Optional[str] = None

@dataclass
class Question:
    id: int
    competencia: str = ""
    evidencia: str = ""
    contenido: str = ""
    contexto: str = ""
    respuesta_correcta: str = ""
    enunciado_parts: List[str] = field(default_factory=list)
    enunciado_imgs: List[str] = field(default_factory=list)
    opciones: List[Opcion] = field(default_factory=list)
    opciones_imgs: Dict[str, str] = field(default_factory=dict)
    ayuda_1_parts: List[str] = field(default_factory=list)
    ayuda_2_parts: List[str] = field(default_factory=list)
    _opciones_text_buffer: List[str] = field(default_factory=list)

    def add_enunciado_text(self, t: str):
        t = sanitize_bullets(t)
        if t:
            self.enunciado_parts.append(t)

    def add_ayuda1_text(self, t: str):
        t = sanitize_bullets(t)
        if t:
            self.ayuda_1_parts.append(t)

    def add_ayuda2_text(self, t: str):
        t = sanitize_bullets(t)
        if t:
            self.ayuda_2_parts.append(t)

    def add_enunciado_image(self, file_rel: str):
        self.enunciado_imgs.append(file_rel)

    def add_opcion_text(self, text: str):
        t = norm(text)
        if not t:
            return
        m = re.match(r"^\s*([A-F])[.)]?\s*(.+)$", t, flags=re.IGNORECASE | re.DOTALL)
        if m and len(m.group(1)) == 1 and len(m.group(2).strip()) > 0:
            self.opciones.append(Opcion(id=m.group(1).upper(), tipo="texto", texto=m.group(2).strip()))
        else:
            self._opciones_text_buffer.append(t)

    def add_opcion_image(self, filename: str):
        next_id = OPTION_LABELS[len(self.opciones)] if len(self.opciones) < len(OPTION_LABELS) else str(len(self.opciones)+1)
        self.opciones.append(Opcion(id=next_id, tipo="imagen", imagen=filename))
        self.opciones_imgs[next_id] = filename

    def finalize_options(self):
        used_ids = set(o.id for o in self.opciones)
        for t in self._opciones_text_buffer:
            next_letter = None
            for L in OPTION_LABELS:
                if L not in used_ids:
                    next_letter = L
                    break
            if next_letter is None:
                next_letter = str(len(self.opciones) + 1)
            self.opciones.append(Opcion(id=next_letter, tipo="texto", texto=t))
            used_ids.add(next_letter)

        def _key(o: Opcion):
            if o.id in OPTION_LABELS:
                return OPTION_LABELS.index(o.id)
            return 999
        self.opciones.sort(key=_key)

    def to_markdown(self, out_dir_images: str, source_file: str) -> str:
        enunciado_md = "\n".join(self.enunciado_parts).strip()
        ayuda1_md = "\n".join(self.ayuda_1_parts).strip()
        ayuda2_md = "\n".join(self.ayuda_2_parts).strip()

        def yescape(s: str) -> str:
            return (s or "").replace('"', '\\"')

        yaml_lines = []
        yaml_lines.append("---")
        yaml_lines.append(f"id: {self.id}")
        yaml_lines.append(f'competencia: "{yescape(self.competencia)}"')
        yaml_lines.append(f'evidencia: "{yescape(self.evidencia)}"')
        yaml_lines.append(f'contenido: "{yescape(self.contenido)}"')
        yaml_lines.append(f'contexto: "{yescape(self.contexto)}"')
        yaml_lines.append(f'respuesta_correcta: "{yescape(self.respuesta_correcta)}"')
        yaml_lines.append("enunciado_md: |")
        for line in (enunciado_md.splitlines() or [""]):
            yaml_lines.append(f"  {line}")
        if self.enunciado_imgs:
            yaml_lines.append("imagenes_enunciado:")
            for p in self.enunciado_imgs:
                yaml_lines.append(f'  - "{p}"')
        else:
            yaml_lines.append("imagenes_enunciado: []")
        yaml_lines.append("opciones:")
        for o in self.opciones:
            yaml_lines.append(f'  - id: "{o.id}"')
            yaml_lines.append(f'    tipo: "{o.tipo}"')
            if o.tipo == "texto":
                yaml_lines.append(f'    texto: "{yescape(o.texto or "")}"')
                yaml_lines.append(f'    imagen: null')
            else:
                yaml_lines.append(f'    texto: null')
                yaml_lines.append(f'    imagen: "{o.imagen}"')
        if self.opciones_imgs:
            yaml_lines.append("imagenes_opciones:")
            for k in sorted(self.opciones_imgs.keys()):
                yaml_lines.append(f'  - id: "{k}"')
                yaml_lines.append(f'    file: "{self.opciones_imgs[k]}"')
        else:
            yaml_lines.append("imagenes_opciones: []")
        yaml_lines.append("ayuda_1_md: |")
        for line in (ayuda1_md.splitlines() or [""]):
            yaml_lines.append(f"  {line}")
        yaml_lines.append("ayuda_2_md: |")
        for line in (ayuda2_md.splitlines() or [""]):
            yaml_lines.append(f"  {line}")
        yaml_lines.append("origen:")
        yaml_lines.append(f'  archivo: "{yescape(os.path.basename(source_file))}"')
        yaml_lines.append(f"  pagina_aproximada: null")
        yaml_lines.append(f"  hash_fuente: null")
        yaml_lines.append("---")
        yaml_lines.append("")
        yaml_lines.append("# Enunciado")
        yaml_lines.append("")
        if enunciado_md:
            yaml_lines.append(enunciado_md)
            yaml_lines.append("")
        for img in self.enunciado_imgs:
            yaml_lines.append(f"![Enunciado imagen]({img})")
        yaml_lines.append("")
        yaml_lines.append("## Opciones")
        for o in self.opciones:
            if o.tipo == "texto":
                yaml_lines.append(f"- {o.id}) {o.texto}")
            else:
                yaml_lines.append(f"- {o.id}) ![Opción {o.id}]({o.imagen})")
        yaml_lines.append("")
        if ayuda1_md:
            yaml_lines.append("## Ayuda 1")
            yaml_lines.append(ayuda1_md)
            yaml_lines.append("")
        if ayuda2_md:
            yaml_lines.append("## Ayuda 2")
            yaml_lines.append(ayuda2_md)
            yaml_lines.append("")
        return "\n".join(yaml_lines)

# -----------------------
# Procesamiento DOCX
# -----------------------

def parse_docx_to_questions(docx_path: str, out_dir: str = "out") -> List['Question']:
    os.makedirs(out_dir, exist_ok=True)
    out_images = os.path.join(out_dir, "images")
    out_questions = os.path.join(out_dir, "questions")
    os.makedirs(out_images, exist_ok=True)
    os.makedirs(out_questions, exist_ok=True)

    doc = Document(docx_path)

    # Mapa id->Question para asegurar que escribimos TODAS las preguntas encontradas
    qmap: Dict[int, Question] = {}

    for block in iter_block_items(doc):
        if isinstance(block, Table):
            created = _parse_table_as_questions(block, out_dir, out_images)
            for q in created:
                qmap[q.id] = q  # si se repite id, la última vista gana (normal)
        # (Si hubiera preguntas fuera de tablas, aquí se podrían manejar)

    # Escribir todas las preguntas
    written = 0
    for qid in sorted(qmap.keys()):
        q = qmap[qid]
        q.finalize_options()
        _write_question(q, out_questions, out_dir, docx_path)
        written += 1

    # Retornar en orden por si se quiere usar programáticamente
    return [qmap[k] for k in sorted(qmap.keys())]

# -----------------------
# Parse de tabla (varias preguntas por tabla)
# -----------------------

def _parse_table_as_questions(table: Table, out_dir: str, out_images: str) -> List[Question]:
    created: List[Question] = []
    q: Optional[Question] = None
    current_section: Optional[str] = None
    enun_img_count = 0

    def push_current():
        nonlocal q
        if q:
            q.finalize_options()
            created.append(q)
            q = None

    def ensure_q(num: Optional[int]):
        nonlocal q, enun_img_count, current_section
        if num is None:
            return
        # si había una pregunta abierta, guárdala antes de abrir otra
        if q:
            push_current()
        q = Question(id=int(num))
        current_section = None
        enun_img_count = 0

    nrows = len(table.rows)
    for ri, row in enumerate(table.rows):
        cells = row.cells
        ncols = len(cells)

        # etiqueta en col 0 (aun cuando la fila tenga una sola celda)
        label = header_key(cell_text(cells[0])) if ncols >= 1 else None

        if label:
            # Cabecera de sección
            if label == "pregunta":
                # buscar número en col 1 o en siguientes 3 filas/celdas
                num = None
                if ncols >= 2:
                    num = looks_like_question_number(cell_text(cells[1]))
                if num is None:
                    look_ahead = min(3, nrows - ri - 1)
                    found = False
                    for k in range(1, look_ahead + 1):
                        for c in table.rows[ri + k].cells:
                            num = looks_like_question_number(cell_text(c))
                            if num is not None:
                                found = True
                                break
                        if found:
                            break
                ensure_q(num)
                continue

            # activar sección y cargar posible contenido en misma fila (col 2..)
            current_section = label
            if q and ncols >= 2:
                content = _join_cells_content(cells[1:])
                _append_content(q, current_section, content)
                for j in range(1, ncols):
                    imgs = extract_images_from_cell(cells[j])
                    for blob, ext in imgs:
                        if current_section == "opciones":
                            letter = _next_option_letter(q)
                            fname = f"q{q.id:03d}-opcion-{letter}{ext}"
                            _save_image(blob, out_images, fname)
                            rel = os.path.relpath(os.path.join(out_images, fname), out_dir).replace("\\", "/")
                            q.add_opcion_image(rel)
                        else:
                            enun_img_count += 1
                            fname = f"q{q.id:03d}-enunciado-{enun_img_count}{ext}"
                            _save_image(blob, out_images, fname)
                            rel = os.path.relpath(os.path.join(out_images, fname), out_dir).replace("\\", "/")
                            q.add_enunciado_image(rel)
            continue

        # Fila sin cabecera: es contenido de la sección vigente
        if q and current_section:
            content = _join_cells_content(cells)
            _append_content(q, current_section, content)
            for j in range(ncols):
                imgs = extract_images_from_cell(cells[j])
                for blob, ext in imgs:
                    if current_section == "opciones":
                        letter = _next_option_letter(q)
                        fname = f"q{q.id:03d}-opcion-{letter}{ext}"
                        _save_image(blob, out_images, fname)
                        rel = os.path.relpath(os.path.join(out_images, fname), out_dir).replace("\\", "/")
                        q.add_opcion_image(rel)
                    else:
                        enun_img_count += 1
                        fname = f"q{q.id:03d}-enunciado-{enun_img_count}{ext}"
                        _save_image(blob, out_images, fname)
                        rel = os.path.relpath(os.path.join(out_images, fname), out_dir).replace("\\", "/")
                        q.add_enunciado_image(rel)

    # cerrar última pregunta de la tabla
    push_current()
    return created

def _join_cells_content(cells: List[_Cell]) -> str:
    parts: List[str] = []
    for c in cells:
        t = cell_text(c)
        if t:
            parts.append(t)
        md_nested = nested_tables_to_md(c)
        if md_nested:
            parts.append(md_nested)
    return "\n".join([p for p in parts if p]).strip()

def _append_content(q: Question, section: str, content: str):
    if not content:
        return
    if section == "competencia":
        q.competencia = (q.competencia + "\n" + content).strip() if q.competencia else content
    elif section == "evidencia":
        q.evidencia = (q.evidencia + "\n" + content).strip() if q.evidencia else content
    elif section == "contenido":
        q.contenido = (q.contenido + "\n" + content).strip() if q.contenido else content
    elif section == "contexto":
        q.contexto = (q.contexto + "\n" + content).strip() if q.contexto else content
    elif section == "respuesta_correcta":
        m = re.search(r"([A-F])", content, flags=re.IGNORECASE)
        q.respuesta_correcta = (m.group(1).upper() if m else norm(content))
    elif section == "enunciado":
        q.add_enunciado_text(content)
    elif section == "opciones":
        for line in content.splitlines():
            line = norm(line)
            if not line:
                continue
            q.add_opcion_text(line)
    elif section == "ayuda_1":
        q.add_ayuda1_text(content)
    elif section == "ayuda_2":
        q.add_ayuda2_text(content)

def _next_option_letter(q: Question) -> str:
    idx = len(q.opciones)
    return OPTION_LABELS[idx] if idx < len(OPTION_LABELS) else str(idx+1)

def _save_image(blob: bytes, out_images: str, filename: str):
    fpath = os.path.join(out_images, filename)
    with open(fpath, "wb") as f:
        f.write(blob)

def _write_question(q: Question, out_questions_dir: str, out_dir: str, source_file: str):
    md = q.to_markdown(out_dir_images=os.path.join(out_dir, "images"), source_file=source_file)
    fname = os.path.join(out_questions_dir, f"q{q.id:03d}.md")
    with open(fname, "w", encoding="utf-8") as f:
        f.write(md)

# -----------------------
# CLI
# -----------------------

def main():
    if len(sys.argv) < 2:
        print("Uso: python scripts/docx_a_markdown.py '<ruta_docx>' [<out_dir>]")
        sys.exit(1)
    docx_path = sys.argv[1]
    out_dir = sys.argv[2] if len(sys.argv) >= 3 else "out"
    questions = parse_docx_to_questions(docx_path, out_dir=out_dir)
    nums = sorted(q.id for q in questions)
    print(f"Listo. Preguntas detectadas: {len(nums)}")
    if nums:
        print(f"Rango detectado: {nums[0]}–{nums[-1]}")
        faltantes = [n for n in range(nums[0], nums[-1]+1) if n not in set(nums)]
        if faltantes:
            print(f"Atención: faltan archivos para: {faltantes}")
    print(f"- Markdown: {os.path.join(out_dir, 'questions')}")
    print(f"- Imágenes: {os.path.join(out_dir, 'images')}")

if __name__ == "__main__":
    main()
